from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "VolunteerHub_Serve_User_Guide.docx"

# compact_reference_guide preset with a restrained VolunteerHub teal override.
NAVY = "16324F"
TEAL = "0F766E"
TEAL_DARK = "115E59"
BLUE = "2E74B5"
MUTED = "667085"
INK = "172033"
LIGHT_TEAL = "E8F5F3"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_run(run, size=11, color=INK, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL_DARK, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run("VOLUNTEERHUB  |  SERVE USER GUIDE"), 8.5, MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Serve Module  |  "), 8.5, MUTED)
    add_page_number(footer)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, 25, NAVY, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        set_run(p.add_run(subtitle), 12.5, MUTED, italic=True)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), 11, INK, bold=True)
        set_run(p.add_run(text[len(bold_lead):]), 11, INK)
    else:
        set_run(p.add_run(text), 11, INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text), 11, INK)
    return p


def begin_numbered_list(doc):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    source_num = next(
        node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_id = int(source_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    next_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_num_id


def add_number(doc, lead, detail, num_id):
    p = doc.add_paragraph(style="List Number")
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
    num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))
    set_run(p.add_run(lead), 11, INK, bold=True)
    set_run(p.add_run(detail), 11, INK)
    return p


def add_callout(doc, label, text, tone="teal"):
    fill = LIGHT_TEAL if tone == "teal" else LIGHT_BLUE if tone == "blue" else LIGHT_GRAY
    label_color = TEAL_DARK if tone == "teal" else NAVY
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(f"{label}: "), 10.5, label_color, bold=True)
    set_run(p.add_run(text), 10.5, INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), 9.5, WHITE, bold=True)
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            cell = table.cell(row_index, col)
            if row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), 9.5, INK, bold=(col == 0))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def new_page(doc, kicker, title, subtitle=None):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(kicker.upper()), 9, TEAL, bold=True)
    add_title(doc, title, subtitle)


def build():
    doc = Document()
    configure_document(doc)

    # Editorial cover.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("VOLUNTEERHUB"), 11, TEAL, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("Serve Module"), 32, NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    set_run(p.add_run("User Guide and Presentation Walkthrough"), 16, TEAL_DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    set_run(
        p.add_run("From finding an opportunity to joining a team and checking in"),
        11.5,
        MUTED,
        italic=True,
    )
    add_callout(
        doc,
        "Audience",
        "Volunteers and presenters demonstrating the volunteer serving journey.",
        "blue",
    )
    add_table(
        doc,
        ["Document", "Details"],
        [
            ("Version", "1.0"),
            ("Implementation reviewed", "June 20, 2026"),
            ("Scope", "Serve module and its volunteer-facing submodules"),
        ],
        [2700, 6660],
    )

    new_page(doc, "Orientation", "Serve at a glance", "A simple path from opportunity to attendance")
    add_callout(
        doc,
        "Purpose",
        "The Serve module helps an approved volunteer discover active upcoming events, choose an Event Team, and request or confirm a place to serve.",
    )
    add_heading(doc, "The six submodules", 1)
    add_table(
        doc,
        ["Submodule", "What the user does"],
        [
            ("1. Browse & Search", "View active upcoming opportunities by month and search by event, team, campus, description, or address."),
            ("2. Event Details", "Open an event to review its date, time, location, description, and available Event Teams."),
            ("3. Directions", "Launch Google Maps using the event coordinates or address."),
            ("4. Team Selection", "Compare Event Team descriptions and see how many positions are filled."),
            ("5. Volunteer & Withdraw", "Join a team, receive a signup result, or leave an active assignment."),
            ("6. Commitment Handoff", "Use My Commitments for directions, leader contact, group chat, and check-in."),
        ],
        [2500, 6860],
    )
    add_heading(doc, "Who can use it", 2)
    add_bullet(doc, "Volunteers use Serve to join their own Event Teams.")
    add_bullet(doc, "The account must be active and the volunteer application must be approved before signup succeeds.")
    add_bullet(doc, "Only active, upcoming events and active Event Teams appear in the volunteer Serve view.")
    add_callout(doc, "Presentation cue", "Describe Serve as the discovery and signup front door. My Commitments is the follow-through workspace after signup.", "blue")

    new_page(doc, "Submodule 1", "Browse and search opportunities", "Find the right event before choosing a team")
    add_heading(doc, "Open the Serve module", 1)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Sign in. ", "Use an active VolunteerHub volunteer account.", numbers)
    add_number(doc, "Select Home. ", "Home opens the Serve experience and displays “Find a place to serve.”", numbers)
    add_number(doc, "Review upcoming months. ", "Event cards are grouped by month and ordered by start time.", numbers)
    add_heading(doc, "Read an event card", 2)
    add_bullet(doc, "Date tile: the calendar date and month.")
    add_bullet(doc, "Team count: the number of active Event Teams within the event.")
    add_bullet(doc, "Event name, time range, and campus or event location.")
    add_bullet(doc, "Directions and Volunteer Here actions.")
    add_heading(doc, "Search effectively", 2)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Enter at least two characters. ", "The app begins a server search after a short typing pause.", numbers)
    add_number(doc, "Use natural terms. ", "Search can match the event name, description, address, campus name, or Event Team name.", numbers)
    add_number(doc, "Clear the search. ", "Return to the complete list of upcoming opportunities.", numbers)
    add_callout(doc, "If nothing appears", "The event may not be active, may already have ended, or may not match the search text. Clear the search before escalating the issue.", "blue")

    new_page(doc, "Submodules 2–4", "Open an event and choose a team", "Review details, get directions, and compare staffing")
    add_heading(doc, "Open event details", 1)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Select Volunteer Here. ", "A details drawer opens without leaving the opportunity list.", numbers)
    add_number(doc, "Confirm the event. ", "Review its campus, description, date, time, and address.", numbers)
    add_number(doc, "Review Available event teams. ", "Each team card contains its name, description, and staffing progress.", numbers)
    add_heading(doc, "Understand staffing", 2)
    add_body(doc, "The filled indicator is displayed as confirmed volunteers divided by required volunteers. For example, 3/5 filled means two confirmed positions remain.")
    add_callout(doc, "Important", "A full team can still display a Volunteer button. The result depends on the team’s signup policy: it may create a waitlist entry or an approval request.", "blue")
    add_heading(doc, "Get directions", 2)
    add_bullet(doc, "Select Directions from the event card to open Google Maps immediately.")
    add_bullet(doc, "VolunteerHub uses saved coordinates when available; otherwise it uses the event or campus address.")
    add_bullet(doc, "Directions open in a new browser tab or the device’s supported maps experience.")
    add_heading(doc, "Choose with confidence", 2)
    add_bullet(doc, "Read the team description before volunteering.")
    add_bullet(doc, "Use staffing progress to identify teams that still need help.")
    add_bullet(doc, "Confirm the event date and location before selecting a team.")

    new_page(doc, "Submodule 5", "Volunteer, status, and withdrawal", "Know what happens after selecting Volunteer")
    add_heading(doc, "Join an Event Team", 1)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Select Volunteer on the desired team card. ", "The request is submitted immediately.", numbers)
    add_number(doc, "Read the confirmation message. ", "Automatic signup with space available confirms the assignment. Other outcomes are sent for follow-up.", numbers)
    add_number(doc, "Open My Commitments. ", "Use it to review the current assignment state and next actions.", numbers)
    add_heading(doc, "Status guide", 2)
    add_table(
        doc,
        ["Status", "Meaning", "What to do"],
        [
            ("Confirmed", "A serving position is reserved.", "Plan to attend; use My Commitments for directions and check-in."),
            ("Requested", "A leader must approve the signup.", "Wait for a decision and review the status later."),
            ("Waitlisted", "Automatic signup found the team at capacity.", "Watch for a later update; a position is not yet reserved."),
            ("Cancelled", "The assignment was withdrawn.", "Choose another team if you still want to serve."),
        ],
        [1800, 3300, 4260],
    )
    add_heading(doc, "Withdraw from a team", 2)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Reopen the event. ", "The team card changes from Volunteer to Withdraw when an active assignment exists.", numbers)
    add_number(doc, "Select Withdraw. ", "VolunteerHub cancels the assignment immediately and refreshes the event.", numbers)
    add_callout(doc, "Before withdrawing", "Because the current action does not display a confirmation prompt, make sure the correct Event Team is selected before pressing Withdraw.", "blue")

    new_page(doc, "Submodule 6", "My Commitments and check-in", "Complete the serving journey after signup")
    add_heading(doc, "Use My Commitments", 1)
    add_body(doc, "Commitments are grouped by month. Each row shows the Event Team, event, time range, location, and current assignment or attendance state.")
    add_table(
        doc,
        ["Action", "When to use it"],
        [
            ("Group chat", "Coordinate with confirmed teammates and leaders. It becomes available after confirmation; team leaders also have access."),
            ("Message leader", "Send a private message through VolunteerHub’s messaging relay."),
            ("Directions", "Open the event destination in Google Maps."),
            ("Check-in", "Record attendance for a confirmed assignment during the allowed time and location window."),
        ],
        [2400, 6960],
    )
    add_heading(doc, "Check in", 2)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Open My Commitments near the event start time. ", "The Check-in action appears for a confirmed assignment.", numbers)
    add_number(doc, "Select Check-in. ", "If location-based self-check-in is enabled, allow the device to share its current location.", numbers)
    add_number(doc, "Wait for “Check-in complete.” ", "The action changes to a Checked in status.", numbers)
    add_callout(doc, "Typical limits", "The default self-check-in window is 30 minutes before through 30 minutes after the event start, within 300 meters. Team configuration can change these limits.", "blue")
    add_heading(doc, "Common check-in messages", 2)
    add_bullet(doc, "Outside the check-in window: try again during the configured arrival period.")
    add_bullet(doc, "Location access is required: enable location permission and retry.")
    add_bullet(doc, "Outside the check-in area: move closer to the event location and retry.")

    new_page(doc, "Presentation aid", "Recommended live demo", "A concise 4–6 minute walkthrough")
    add_heading(doc, "Demo sequence", 1)
    numbers = begin_numbered_list(doc)
    add_number(doc, "Set the story. ", "“A volunteer wants to help at an upcoming service and needs to find the team with the greatest need.”", numbers)
    add_number(doc, "Show Home / Serve. ", "Point out monthly grouping, the event card, and the team count.", numbers)
    add_number(doc, "Demonstrate search. ", "Search with a campus, event, or Event Team name.", numbers)
    add_number(doc, "Open Volunteer Here. ", "Review date, location, directions, team descriptions, and staffing progress.", numbers)
    add_number(doc, "Select Volunteer. ", "Explain Confirmed, Requested, and Waitlisted outcomes.", numbers)
    add_number(doc, "Open My Commitments. ", "Show group chat, message leader, directions, and the check-in handoff.", numbers)
    add_heading(doc, "Presenter talking points", 2)
    add_bullet(doc, "Serve is focused: volunteers see active, upcoming opportunities rather than administrative event setup.")
    add_bullet(doc, "The volunteer chooses an Event Team, not just a general event.")
    add_bullet(doc, "Staffing visibility helps volunteers place themselves where help is needed.")
    add_bullet(doc, "Signup policy controls whether the result is immediate confirmation, leader approval, or a waitlist.")
    add_bullet(doc, "My Commitments carries the volunteer from signup through communication, directions, and attendance.")
    add_callout(doc, "Demo preparation", "Use an approved volunteer account and an active future event with at least one active Event Team. For a clean check-in demo, configure the event near the presentation time and location.", "teal")

    new_page(doc, "Quick reference", "Troubleshooting and presenter Q&A")
    add_table(
        doc,
        ["Question or issue", "Recommended answer"],
        [
            ("Why can’t I see an event?", "Volunteer Serve shows only active events that have not ended. Clear search text and confirm the event is active."),
            ("Why can’t I volunteer?", "The volunteer profile must be active and the application approved. The event and team must also be active."),
            ("Why was I not confirmed?", "The team may require leader approval or may already be full, producing Requested or Waitlisted status."),
            ("Can I join twice?", "No. VolunteerHub prevents more than one active assignment for the same Event Team."),
            ("Can I change teams?", "Withdraw from the current team, then volunteer for the preferred team. A dedicated move/swap flow is not currently exposed."),
            ("Why does check-in fail?", "Confirm the assignment, time window, device location permission, and distance from the event."),
            ("Does leaving a team promote the next person?", "Automatic waitlist promotion is not currently implemented; a leader may need to follow up."),
        ],
        [3100, 6260],
    )
    add_heading(doc, "One-sentence summary", 1)
    add_callout(doc, "Serve", "Discover an active event, choose the right Event Team, volunteer with a clear status, and carry the commitment through directions, communication, and check-in.", "teal")

    doc.core_properties.title = "VolunteerHub Serve Module User Guide"
    doc.core_properties.subject = "Volunteer-facing Serve module and submodules"
    doc.core_properties.author = "VolunteerHub"
    doc.core_properties.keywords = "VolunteerHub, Serve, volunteer, Event Team, commitments, check-in"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
