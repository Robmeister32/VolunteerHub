from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MODULE_DOCUMENTATION.md"
OUTPUT = ROOT / "VolunteerHub_Module_Documentation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "20384F"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "111827"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[index])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VOLUNTEERHUB")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Module Documentation")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    run = p.add_run("Functional specifications, implemented business rules, and integration boundaries")
    run.italic = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.75, 4.75])
    set_repeat_table_header(table.rows[0])
    values = (
        ("Document version", "1.0"),
        ("Implementation reviewed", "June 15, 2026"),
        ("Source of truth", "Express API, React clients, PostgreSQL migrations, views, and backend tests"),
    )
    for row, values_row in zip(table.rows, values):
        row.cells[0].text = values_row[0]
        row.cells[1].text = values_row[1]
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_page_break()


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run("VolunteerHub Module Documentation  |  ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(p)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("VolunteerHub  |  Application Reference")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    width_patterns = {
        2: [1.85, 4.65],
        3: [1.55, 1.2, 3.75],
        4: [1.35, 1.3, 1.45, 2.4],
    }
    set_table_widths(table, width_patterns.get(columns, [6.5 / columns] * columns))
    for row_index, row in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            add_inline(p, row[col_index] if col_index < len(row) else "")
            for run in p.runs:
                run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)
    add_header(section)
    add_footer(section)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            index += 1
            continue
        p = doc.add_paragraph()
        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        else:
            add_inline(p, line)
        index += 1

    properties = doc.core_properties
    properties.title = "VolunteerHub Module Documentation"
    properties.subject = "Functional specifications and business rules"
    properties.author = "VolunteerHub"
    properties.keywords = "VolunteerHub, modules, specifications, business rules"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
